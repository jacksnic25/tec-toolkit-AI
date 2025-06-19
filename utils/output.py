from enum import Enum
from dataclasses import dataclass
from typing import Optional, Dict, Any
from rich.console import Console
from rich.table import Table
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches
import os

console = Console()

class Status(Enum):
    SUCCESS = 1
    ERROR = 2

@dataclass
class Result:
    status: Status
    data: Optional[Dict[str, Any]] = None
    message: Optional[str] = None

def format_vt_url_report(data: dict) -> str:
    # existing code unchanged
    pass

def display_results(result: Result):
    # existing code unchanged
    pass

try:
    from weasyprint import HTML
except ImportError:
    HTML = None

def generate_html_report(data: dict) -> str:
    # Generate a simple HTML report wrapping JSON string in <pre> tags
    import json
    if isinstance(data, str):
        try:
            json_obj = json.loads(data)
            pretty_json = json.dumps(json_obj, indent=4)
        except Exception:
            pretty_json = data
    else:
        pretty_json = json.dumps(data, indent=4)
    html_content = f"""
    <html>
    <head><title>Report</title></head>
    <body>
    <h1>Report</h1>
    <pre>{pretty_json}</pre>
    </body>
    </html>
    """
    return html_content

def generate_vt_url_html_report(data: dict) -> str:
    # existing code unchanged
    pass

def generate_vt_file_html_report(data: dict) -> str:
    # existing code unchanged
    pass

def generate_vt_file_docx_report(data: dict, file_path: str):
    """
    Generate a DOCX report for VirusTotal file scan results with columns and rows.
    """
    document = Document()
    document.add_heading('VirusTotal File Scan Report', 0)

    sha256 = data.get("sha256", "N/A")
    malicious = data.get("malicious", "N/A")
    document.add_paragraph(f"SHA256: {sha256}")
    document.add_paragraph(f"Malicious Detections: {malicious}")

    analysis_results = data.get("analysis", {})
    if not analysis_results:
        document.add_paragraph("No analysis data available.")
        document.save(file_path)
        return

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Engine'
    hdr_cells[1].text = 'Category'
    hdr_cells[2].text = 'Result'

    for engine_name, result_data in analysis_results.items():
        row_cells = table.add_row().cells
        row_cells[0].text = engine_name
        row_cells[1].text = result_data.get("category", "unknown")
        row_cells[2].text = result_data.get("result", "N/A")

    document.save(file_path)

def save_report(result: Result, file_path: str, file_format: str):
    """
    Save the result report to a file in the specified format.
    Supported formats: json, html, pdf, doc
    """
    if result.status != Status.SUCCESS:
        raise ValueError(f"Cannot save report: {result.message}")

    if not file_path:
        raise ValueError("Output file path must be specified for report saving.")

    file_format = file_format.lower()
    data = result.data
    # If data is a string and looks like JSON, try to parse it
    import json
    try:
        parsed_data = json.loads(data) if isinstance(data, str) else data
    except Exception:
        parsed_data = data

    if file_format == "json":
        # Save as pretty JSON
        with open(file_path, "w", encoding="utf-8") as f:
            json.dump(parsed_data, f, indent=4)
    elif file_format == "html":
        # Use special HTML report for VirusTotal URL scan results
        if isinstance(parsed_data, dict) and "url" in parsed_data and "analysis" in parsed_data:
            html_content = generate_vt_url_html_report(parsed_data)
        elif isinstance(parsed_data, dict) and "sha256" in parsed_data and "analysis" in parsed_data:
            html_content = generate_vt_file_html_report(parsed_data)
        else:
            html_content = generate_html_report(parsed_data)
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(html_content)
    elif file_format == "pdf":
        if HTML is None:
            raise ImportError("WeasyPrint is not installed. Please install it to enable PDF output.")
        # Use special HTML report for VirusTotal URL scan results
        if isinstance(parsed_data, dict) and "url" in parsed_data and "analysis" in parsed_data:
            html_content = generate_vt_url_html_report(parsed_data)
        elif isinstance(parsed_data, dict) and "sha256" in parsed_data and "analysis" in parsed_data:
            html_content = generate_vt_file_html_report(parsed_data)
        else:
            html_content = generate_html_report(parsed_data)
        HTML(string=html_content).write_pdf(file_path)
    elif file_format == "doc":
        if not file_path.lower().endswith(".docx"):
            file_path += ".docx"
        if isinstance(parsed_data, dict) and "sha256" in parsed_data and "analysis" in parsed_data:
            generate_vt_file_docx_report(parsed_data, file_path)
        else:
            # For other data, create a simple doc with JSON text
            document = Document()
            document.add_heading('Report', 0)
            document.add_paragraph(json.dumps(parsed_data, indent=4))
            document.save(file_path)
    else:
        raise ValueError(f"Unsupported report format: {file_format}")
